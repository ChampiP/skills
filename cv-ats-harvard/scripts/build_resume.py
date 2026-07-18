#!/usr/bin/env python3
"""
Genera un CV en formato Harvard/ATS (.docx) a partir de un JSON de datos
estructurados, reutilizando los estilos (Heading 1, Body Text, List Paragraph,
Normal) de la plantilla Harvard oficial.

Uso:
    python build_resume.py --data datos.json --template assets/harvard_template.docx \
        --out /mnt/user-data/outputs/CV_Nombre_ATS.docx

El JSON debe seguir el esquema documentado en SKILL.md /
references/harvard_sections.md.
"""

import argparse
import json
import sys
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Posición fija de la tabulación para las líneas "texto izquierda / fecha-lugar derecha".
# Se usa una tabulación IZQUIERDA (no RIGHT) en una posición fija, igual que la plantilla
# oficial de Harvard FAS -- una tabulación RIGHT con leader de espacios es lo que produce
# la superposición de texto que se ve al abrir el documento en Word/Google Docs.
DATE_TAB_POSITION = Inches(6.3)

# Fuente ATS-safe pero más moderna que Times New Roman (que trae la plantilla original).
# Arial/Calibri/Helvetica son igual de válidas para ATS; se deja configurable por si
# el usuario prefiere otra de la lista permitida.
FONT_NAME = "Calibri"


def set_document_font(doc, font_name=FONT_NAME):
    """Sustituye la fuente (Times New Roman) de los estilos base por una más moderna,
    sin tocar nada de la estructura ATS (sigue siendo texto plano, una columna)."""
    for style_name in ("Normal", "Heading 1", "Body Text", "List Paragraph"):
        style = doc.styles[style_name]
        style.font.name = font_name
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rfonts.set(qn(attr), font_name)


def add_bottom_border(paragraph):
    """Agrega una línea horizontal simple bajo el párrafo (borde de párrafo nativo de
    Word, no una tabla ni un elemento gráfico -- no afecta la lectura ATS)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_two_col_line(doc, left, right, style="Normal"):
    """Línea con texto a la izquierda y fecha/lugar tabulado, replicando el patrón
    de la plantilla Harvard (tabulación izquierda fija, no alineación derecha real)."""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.tab_stops.add_tab_stop(DATE_TAB_POSITION, WD_TAB_ALIGNMENT.LEFT)
    p.add_run(f"{left}\t{right}")
    return p


def add_section_heading(doc, text):
    """Título de sección centrado, igual que 'Education' / 'Experience' en la plantilla."""
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Paragraph")


def build(data, template_path, out_path):
    doc = Document(template_path)
    set_document_font(doc)

    # Limpia todo el contenido de ejemplo de la plantilla, conservando estilos.
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

    # 1. Encabezado -- centrado, igual que "Firstname Lastname" en la plantilla Harvard.
    p_nombre = doc.add_paragraph(style="Normal")
    p_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_nombre = p_nombre.add_run(data["nombre"])
    run_nombre.bold = True
    run_nombre.font.size = Pt(16)

    # 1b. Subtítulo de rol/título profesional (opcional), como en muchos CV modernos.
    if data.get("titulo_profesional"):
        p_rol = doc.add_paragraph(style="Body Text")
        p_rol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_rol = p_rol.add_run(data["titulo_profesional"])
        run_rol.italic = True

    p_contacto = doc.add_paragraph(data["contacto"], style="Body Text")
    p_contacto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(p_contacto)

    # 2. Perfil profesional
    if data.get("perfil"):
        add_section_heading(doc, "Perfil Profesional")
        doc.add_paragraph(data["perfil"], style="Normal")

    # 3. Educación
    if data.get("educacion"):
        add_section_heading(doc, "Educación")
        for e in data["educacion"]:
            add_two_col_line(doc, e["institucion"], e.get("lugar", ""))
            add_two_col_line(doc, e["titulo"], e.get("fecha", ""), style="Body Text")

    # 4. Experiencia laboral
    if data.get("experiencia"):
        add_section_heading(doc, "Experiencia Laboral")
        for exp in data["experiencia"]:
            add_two_col_line(doc, exp["empresa"], exp.get("lugar", ""))
            add_two_col_line(doc, exp["cargo"], exp.get("fecha", ""), style="Body Text")
            add_bullets(doc, exp.get("logros", []))

    # 5. Proyectos relevantes
    if data.get("proyectos"):
        add_section_heading(doc, "Proyectos Relevantes")
        for proj in data["proyectos"]:
            doc.add_paragraph(proj["nombre"], style="Normal")
            add_bullets(doc, proj.get("logros", []))

    # 6. Habilidades
    if data.get("habilidades"):
        add_section_heading(doc, "Habilidades")
        for h in data["habilidades"]:
            p = doc.add_paragraph(style="Body Text")
            p.add_run(f"{h['categoria']}: ").bold = True
            p.add_run(h["items"])

    # 7. Certificaciones
    if data.get("certificaciones"):
        add_section_heading(doc, "Certificaciones")
        for c in data["certificaciones"]:
            linea = f"{c['institucion']} — {c['curso']} — {c['fecha']}"
            doc.add_paragraph(linea, style="List Paragraph")

    # 8. Referencias (opcional)
    if data.get("referencias"):
        add_section_heading(doc, "Referencias")
        for r in data["referencias"]:
            doc.add_paragraph(r["nombre"], style="Normal")
            doc.add_paragraph(f"{r.get('cargo', '')} — {r.get('contacto', '')}", style="Body Text")

    doc.save(out_path)
    print(f"CV generado en: {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", required=True, help="Ruta al JSON con los datos del CV")
    ap.add_argument("--template", required=True, help="Ruta a la plantilla Harvard (.docx)")
    ap.add_argument("--out", required=True, help="Ruta de salida del .docx generado")
    args = ap.parse_args()

    with open(args.data, encoding="utf-8") as f:
        data = json.load(f)

    build(data, args.template, args.out)


if __name__ == "__main__":
    main()
